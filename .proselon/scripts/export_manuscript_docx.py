#!/usr/bin/env python3
"""Export manuscript scene files into a Word (.docx) document.

Usage:
    python .proselon/scripts/export_manuscript_docx.py "Book 1" "Favorite Person" ["Author Name"]

Arguments:
    book_folder:  Name of the folder under Manuscripts/ (e.g. "Book 1")
    book_title:   Title used for metadata and output filename (e.g. "Favorite Person")
    author:       Optional author or pen name for the title page and metadata

Dependencies:
    pip3 install python-docx
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(
        "Error: the python-docx library is not installed.\n"
        "Install it with:  pip3 install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

FONT_NAME = "Georgia"
BODY_FONT_SIZE = Pt(11)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def natural_sort_key(path):
    """Sort by embedded numbers so Chapter 2 comes before Chapter 10."""
    return [
        int(part) if part.isdigit() else part.lower()
        for part in re.split(r"(\d+)", path.name)
    ]


def chapter_heading(dir_name):
    """Heading for a chapter folder.

    "Chapter 0 - Prologue" -> "Prologue", "Chapter 13 - Epilogue" -> "Epilogue",
    "Chapter 3" -> "Chapter 3".
    """
    if " - " in dir_name:
        return dir_name.split(" - ", 1)[1].strip()
    num = re.search(r"\d+", dir_name)
    return f"Chapter {num.group()}" if num else dir_name


def add_formatted_text(paragraph, text):
    """Parse markdown inline formatting (**bold**, *italic*) into Word runs."""
    pattern = re.compile(r"(\*{3}(.+?)\*{3}|\*{2}(.+?)\*{2}|\*(.+?)\*)")

    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.name = FONT_NAME
            run.font.size = BODY_FONT_SIZE

        if match.group(2):       # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):     # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):     # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True

        run.font.name = FONT_NAME
        run.font.size = BODY_FONT_SIZE
        last_end = match.end()

    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = FONT_NAME
        run.font.size = BODY_FONT_SIZE


def add_scene_break(doc):
    """Insert a centred bullet scene-break separator."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\u2022  \u2022  \u2022")
    run.font.name = FONT_NAME
    run.font.size = BODY_FONT_SIZE
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_markdown_block(doc, md_text):
    """Add a block of markdown text, handling headings, lists, and paragraphs."""
    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            for run in p.runs:
                run.font.name = FONT_NAME
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            for run in p.runs:
                run.font.name = FONT_NAME
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, stripped[2:])
        elif stripped == "---":
            add_scene_break(doc)
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, stripped)


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def export(book_folder, book_title, author=None):
    project_root = Path(__file__).resolve().parent.parent.parent
    manuscripts_dir = project_root / "Manuscripts" / book_folder

    if not manuscripts_dir.is_dir():
        print("Error: %s not found." % manuscripts_dir, file=sys.stderr)
        sys.exit(1)

    # --- Discover chapters ------------------------------------------------
    chapter_dirs = sorted(
        [d for d in manuscripts_dir.iterdir() if d.is_dir() and "Chapter" in d.name],
        key=natural_sort_key,
    )
    if not chapter_dirs:
        print("Error: No chapter folders found in %s." % manuscripts_dir, file=sys.stderr)
        sys.exit(1)

    # --- Create document --------------------------------------------------
    doc = Document()

    # Document metadata
    doc.core_properties.title = book_title
    if author:
        doc.core_properties.author = author

    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = BODY_FONT_SIZE
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # --- Title page -------------------------------------------------------
    for _ in range(8):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(book_title)
    run.font.name = FONT_NAME
    run.font.size = Pt(28)
    run.bold = True

    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run("by " + author)
        run.font.name = FONT_NAME
        run.font.size = Pt(14)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # --- Manuscript chapters ----------------------------------------------
    total_scenes = 0

    for chapter_dir in chapter_dirs:
        scene_files = sorted(
            [f for f in chapter_dir.iterdir()
             if f.is_file() and f.suffix == ".md" and f.name.startswith("S")],
            key=natural_sort_key,
        )

        h = doc.add_heading(chapter_heading(chapter_dir.name), level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.name = FONT_NAME

        for i, scene_file in enumerate(scene_files):
            content = scene_file.read_text(encoding="utf-8").strip()
            add_markdown_block(doc, content)
            total_scenes += 1

            if i < len(scene_files) - 1:
                add_scene_break(doc)

        doc.add_page_break()

    # --- Write output -----------------------------------------------------
    output_path = manuscripts_dir.parent / (book_folder + " - " + book_title + ".docx")
    doc.save(str(output_path))

    print("Exported: " + output_path.name)
    print("Chapters: %d, Scenes: %d" % (len(chapter_dirs), total_scenes))
    size_kb = output_path.stat().st_size / 1024
    print("Size: %.0f KB" % size_kb)


if __name__ == "__main__":
    if len(sys.argv) not in (3, 4):
        print(__doc__.strip())
        sys.exit(1)
    export(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) == 4 else None)
